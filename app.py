import streamlit as st
import pandas as pd
import openai
import os
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from io import BytesIO
import json
import re
import plotly.graph_objects as go
import plotly.express as px
from collections import defaultdict
import networkx as nx
import time
from datetime import datetime, timedelta

# Load your API Key
load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=api_key)

# Tyler Technologies brand colors
TYLER_BLUE = "#003F87"
TYLER_LIGHT_BLUE = "#0075C9"
TYLER_GRAY = "#6C757D"
TYLER_DARK_GRAY = "#333333"

# Batch processing configuration
BATCH_SIZE = 50  # Process 50 programs at a time
SAVE_INTERVAL = 10  # Save progress every 10 programs

# Helper functions for visualizations
def analyze_process_overlaps(program_processes):
    """Analyze overlaps between programs based on process types and names"""
    overlap_data = {}
    programs = list(program_processes.keys())
    
    for i, prog1 in enumerate(programs):
        for j, prog2 in enumerate(programs):
            if i < j:  # Only compare each pair once
                processes1 = program_processes[prog1]
                processes2 = program_processes[prog2]
                
                # Find matching process types
                types1 = set(p['process_type'] for p in processes1)
                types2 = set(p['process_type'] for p in processes2)
                common_types = types1.intersection(types2)
                
                # Calculate similarity score
                total_types = len(types1.union(types2))
                similarity_score = len(common_types) / total_types if total_types > 0 else 0
                
                # Find similar processes
                matching_processes = []
                for p1 in processes1:
                    for p2 in processes2:
                        if p1['process_type'] == p2['process_type']:
                            matching_processes.append({
                                'process_type': p1['process_type'],
                                'process1': p1['process_name'],
                                'process2': p2['process_name']
                            })
                
                overlap_data[(prog1, prog2)] = {
                    'common_types': list(common_types),
                    'similarity_score': similarity_score,
                    'matching_processes': matching_processes
                }
    
    return overlap_data

def create_overlap_matrix(overlap_data, program_processes):
    """Create a heatmap showing process overlap between programs"""
    programs = list(program_processes.keys())
    n_programs = len(programs)
    
    # For large datasets, show only the most connected programs
    if n_programs > 50:
        st.warning(f"Showing top 50 most connected programs out of {n_programs} for clarity")
        
        # Calculate connection scores for each program
        connection_scores = {}
        for prog in programs:
            score = 0
            for (p1, p2), data in overlap_data.items():
                if (prog == p1 or prog == p2) and data['similarity_score'] > 0.3:
                    score += data['similarity_score']
            connection_scores[prog] = score
        
        # Get top 50 most connected programs
        top_programs = sorted(connection_scores.items(), key=lambda x: x[1], reverse=True)[:50]
        programs = [p[0] for p in top_programs]
        n_programs = len(programs)
    
    # Create similarity matrix
    matrix = [[0] * n_programs for _ in range(n_programs)]
    
    for i, prog1 in enumerate(programs):
        matrix[i][i] = 1.0  # Self-similarity is 1
        for j, prog2 in enumerate(programs):
            if i < j:
                key = (prog1, prog2)
                if key in overlap_data:
                    similarity = overlap_data[key]['similarity_score']
                    matrix[i][j] = similarity
                    matrix[j][i] = similarity
    
    # Create heatmap with better sizing for large datasets
    height = max(600, min(1200, n_programs * 15))  # Dynamic height based on programs
    
    # Truncate long program names for display
    display_programs = [p[:30] + "..." if len(p) > 30 else p for p in programs]
    
    fig = go.Figure(data=go.Heatmap(
        z=matrix,
        x=display_programs,
        y=display_programs,
        colorscale=[[0, '#FFFFFF'], [0.5, TYLER_LIGHT_BLUE], [1, TYLER_BLUE]],
        text=[[f'{val:.2f}' for val in row] for row in matrix],
        texttemplate='%{text}',
        textfont={"size": 8},
        colorbar=dict(title="Similarity Score"),
        hoverongaps=False
    ))
    
    fig.update_layout(
        title="Process Overlap Matrix (Top Connected Programs)" if n_programs == 50 else "Process Overlap Matrix",
        xaxis_title="Programs",
        yaxis_title="Programs",
        height=height,
        xaxis={'tickangle': -90, 'tickfont': {'size': 10}},
        yaxis={'tickfont': {'size': 10}},
        margin=dict(l=200, r=50, t=50, b=200)  # More margin for labels
    )
    
    st.plotly_chart(fig, use_container_width=True)

def create_process_type_chart(program_processes):
    """Create a bar chart showing distribution of process types"""
    process_type_counts = defaultdict(int)
    
    for program, processes in program_processes.items():
        for process in processes:
            process_type_counts[process['process_type']] += 1
    
    # Convert to dataframe
    df = pd.DataFrame([
        {'Process Type': ptype, 'Count': count}
        for ptype, count in process_type_counts.items()
    ])
    df = df.sort_values('Count', ascending=False)
    
    # Create bar chart
    fig = px.bar(df, x='Process Type', y='Count',
                 color='Count',
                 color_continuous_scale=[TYLER_LIGHT_BLUE, TYLER_BLUE],
                 title="Distribution of Process Types Across Programs")
    
    fig.update_layout(
        xaxis_tickangle=-45,
        height=500,
        showlegend=False
    )
    
    st.plotly_chart(fig, use_container_width=True)

def create_network_graph(overlap_data, program_processes):
    """Create a network graph showing program connections based on process overlap"""
    programs = list(program_processes.keys())
    
    # For large datasets, filter to show only significant connections
    if len(programs) > 100:
        st.warning("Showing only programs with significant overlaps (>0.5 similarity) for clarity")
        min_similarity = 0.5
    else:
        min_similarity = 0.3
    
    # Create network graph
    G = nx.Graph()
    
    # Find programs that have significant connections
    connected_programs = set()
    for (prog1, prog2), data in overlap_data.items():
        if data['similarity_score'] > min_similarity:
            connected_programs.add(prog1)
            connected_programs.add(prog2)
    
    # Add only connected programs as nodes
    for program in connected_programs:
        G.add_node(program)
    
    # Add edges based on similarity
    edge_trace_list = []
    
    for (prog1, prog2), data in overlap_data.items():
        if data['similarity_score'] > min_similarity and prog1 in connected_programs and prog2 in connected_programs:
            G.add_edge(prog1, prog2, weight=data['similarity_score'])
    
    if len(G.nodes()) == 0:
        st.info("No programs with significant overlaps found. Try adjusting the similarity threshold.")
        return
    
    # Use a layout that handles large graphs better
    try:
        if len(G.nodes()) > 50:
            pos = nx.kamada_kawai_layout(G)
        else:
            pos = nx.spring_layout(G, k=3, iterations=50)
    except:
        # Fallback to circular layout if other layouts fail
        pos = nx.circular_layout(G)
    
    # Create edge traces
    for edge in G.edges(data=True):
        x0, y0 = pos[edge[0]]
        x1, y1 = pos[edge[1]]
        weight = edge[2]['weight']
        
        edge_trace = go.Scatter(
            x=[x0, x1, None],
            y=[y0, y1, None],
            line=dict(width=weight*3, color=TYLER_LIGHT_BLUE),
            hoverinfo='none',
            mode='lines',
            showlegend=False
        )
        edge_trace_list.append(edge_trace)
    
    # Create node trace
    node_x = []
    node_y = []
    node_text = []
    node_connections = []
    
    for node in G.nodes():
        x, y = pos[node]
        node_x.append(x)
        node_y.append(y)
        # Truncate long names and add connection count
        truncated_name = node[:25] + "..." if len(node) > 25 else node
        connections = len(list(G.neighbors(node)))
        node_text.append(f"{truncated_name}<br>({connections} connections)")
        node_connections.append(connections)
    
    # Size nodes based on number of connections
    node_sizes = [20 + (conn * 5) for conn in node_connections]
    
    node_trace = go.Scatter(
        x=node_x,
        y=node_y,
        mode='markers+text',
        text=[n.split('<br>')[0] for n in node_text],  # Show only name as label
        hovertext=node_text,  # Show full info on hover
        textposition="top center",
        hoverinfo='text',
        marker=dict(
            size=node_sizes,
            color=node_connections,
            colorscale=[[0, TYLER_LIGHT_BLUE], [1, TYLER_BLUE]],
            colorbar=dict(title="Connections"),
            line_width=2,
            line_color='white'
        ),
        textfont=dict(size=8)
    )
    
    # Create figure
    fig = go.Figure(data=edge_trace_list + [node_trace])
    
    fig.update_layout(
        title=f"Program Process Overlap Network ({len(G.nodes())} programs shown)",
        showlegend=False,
        hovermode='closest',
        margin=dict(b=20,l=5,r=5,t=40),
        xaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
        yaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
        height=800,
        plot_bgcolor='white'
    )
    
    st.plotly_chart(fig, use_container_width=True)
    
    # Add summary statistics
    st.info(f"""
    Network Statistics:
    - Programs shown: {len(G.nodes())} (out of {len(programs)} total)
    - Connections shown: {len(G.edges())}
    - Minimum similarity threshold: {min_similarity}
    - Most connected program: {max(node_connections)} connections
    """)

def create_cross_department_analysis(overlap_data, program_processes):
    """Create analysis of cross-departmental process overlaps"""
    # Identify which programs might be in different departments
    # This is a simplified analysis - in real use, you'd have department info
    
    cross_dept_overlaps = []
    
    for (prog1, prog2), data in overlap_data.items():
        if data['similarity_score'] > 0.5:
            for process_match in data['matching_processes']:
                cross_dept_overlaps.append({
                    'Program 1': prog1,
                    'Program 2': prog2,
                    'Process Type': process_match['process_type'],
                    'Similarity': data['similarity_score']
                })
    
    if cross_dept_overlaps:
        df = pd.DataFrame(cross_dept_overlaps)
        
        # Create grouped bar chart
        fig = px.bar(df, x='Process Type', color='Similarity',
                     title="High-Similarity Process Types Across Programs",
                     color_continuous_scale=[TYLER_LIGHT_BLUE, TYLER_BLUE])
        
        fig.update_layout(
            xaxis_tickangle=-45,
            height=500
        )
        
        st.plotly_chart(fig, use_container_width=True)
        
        # Show detailed table
        st.subheader("Detailed Cross-Program Process Overlaps")
        st.dataframe(df.sort_values('Similarity', ascending=False))
    else:
        st.info("No significant cross-program overlaps detected.")

def create_overlap_matrix_data(overlap_data, program_processes):
    """Create a dataframe for the overlap matrix"""
    programs = list(program_processes.keys())
    matrix_data = {}
    
    for prog in programs:
        matrix_data[prog] = {}
        for other_prog in programs:
            if prog == other_prog:
                matrix_data[prog][other_prog] = 1.0
            else:
                key = (prog, other_prog) if (prog, other_prog) in overlap_data else (other_prog, prog)
                if key in overlap_data:
                    matrix_data[prog][other_prog] = overlap_data[key]['similarity_score']
                else:
                    matrix_data[prog][other_prog] = 0.0
    
    return pd.DataFrame(matrix_data)

def create_word_doc(program_processes, overlap_analysis, analyses, process_overlap_data):
    doc = Document()
    doc.add_heading('Program Process Analysis Report', 0)
    
    # Add process overlap analysis first
    doc.add_heading('Process Overlap Analysis', level=1)
    doc.add_paragraph(overlap_analysis)
    doc.add_page_break()
    
    # Add process summary
    doc.add_heading('Process Summary by Program', level=1)
    for program, processes in program_processes.items():
        doc.add_heading(program, level=2)
        for proc in processes:
            doc.add_paragraph(f"• {proc['process_name']} ({proc['process_type']})")
            doc.add_paragraph(f"  {proc['description']}", style='Normal')
    doc.add_page_break()
    
    # Add individual analyses
    doc.add_heading('Individual Program Efficiency Analyses', level=1)
    for item in analyses:
        doc.add_heading(item['Program'], level=2)
        doc.add_paragraph(item['Analysis'])
        doc.add_page_break()
    
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def create_excel_file(program_processes, overlap_analysis, analyses, process_overlap_data):
    bio = BytesIO()
    with pd.ExcelWriter(bio, engine='openpyxl') as writer:
        # Sheet 1: Process Summary
        process_data = []
        for program, processes in program_processes.items():
            for proc in processes:
                process_data.append({
                    'Program': program,
                    'Process Name': proc['process_name'],
                    'Process Type': proc['process_type'],
                    'Description': proc['description']
                })
        
        if process_data:
            process_df = pd.DataFrame(process_data)
            process_df.to_excel(writer, index=False, sheet_name='Process Summary')
        
        # Sheet 2: Process Overlap Matrix
        overlap_matrix_df = create_overlap_matrix_data(process_overlap_data, program_processes)
        overlap_matrix_df.to_excel(writer, sheet_name='Process Overlap Matrix')
        
        # Sheet 3: Overlap Details
        overlap_details = []
        for (prog1, prog2), similarity_data in process_overlap_data.items():
            if similarity_data['similarity_score'] > 0.3:  # Threshold for relevance
                for match in similarity_data['matching_processes']:
                    overlap_details.append({
                        'Program 1': prog1,
                        'Program 2': prog2,
                        'Process Type': match['process_type'],
                        'Process 1': match['process1'],
                        'Process 2': match['process2'],
                        'Similarity Score': round(similarity_data['similarity_score'], 2),
                        'Consolidation Opportunity': 'High' if similarity_data['similarity_score'] > 0.7 else 'Medium'
                    })
        
        if overlap_details:
            overlap_details_df = pd.DataFrame(overlap_details)
            overlap_details_df.to_excel(writer, index=False, sheet_name='Overlap Details')
        
        # Sheet 4: Overlap Analysis Text
        overlap_df = pd.DataFrame([{'Overlap Analysis': overlap_analysis}])
        overlap_df.to_excel(writer, index=False, sheet_name='Overlap Analysis')
        
        # Sheet 5: Individual Analyses
        analyses_df = pd.DataFrame(analyses)
        analyses_df.to_excel(writer, index=False, sheet_name='Program Analyses')
        
        # Auto-adjust column widths
        for sheet_name in writer.sheets:
            worksheet = writer.sheets[sheet_name]
            for column in worksheet.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 100)
                worksheet.column_dimensions[column_letter].width = adjusted_width
    
    bio.seek(0)
    return bio

def save_progress(progress_data):
    """Save progress to session state"""
    st.session_state.saved_progress = progress_data
    st.session_state.last_save = datetime.now()

def load_progress():
    """Load progress from session state"""
    if 'saved_progress' in st.session_state:
        return st.session_state.saved_progress
    return None

def estimate_time_remaining(processed, total, elapsed_time):
    """Estimate time remaining based on current progress"""
    if processed == 0:
        return "Calculating..."
    
    avg_time_per_item = elapsed_time / processed
    remaining_items = total - processed
    estimated_seconds = avg_time_per_item * remaining_items
    
    return str(timedelta(seconds=int(estimated_seconds)))

# Custom CSS for Tyler Technologies branding
st.markdown("""
<style>
    /* Wider layout */
    .block-container {
        max-width: 1200px;
        padding: 2rem 1rem;
    }
    
    /* Main header styling */
    .main > div {
        padding-top: 2rem;
    }
    
    /* Headers */
    h1 {
        color: #003F87 !important;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
        font-weight: 600;
    }
    
    h2 {
        color: #003F87 !important;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    }
    
    h3 {
        color: #0075C9 !important;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    }
    
    /* Button styling */
    .stButton > button {
        background-color: #003F87;
        color: white;
        border: none;
        padding: 0.5rem 1rem;
        font-weight: 500;
        transition: background-color 0.3s;
    }
    
    .stButton > button:hover {
        background-color: #0075C9;
        color: white;
    }
    
    /* Progress bar */
    .stProgress > div > div > div > div {
        background-color: #0075C9;
    }
    
    /* Expander styling */
    .streamlit-expanderHeader {
        background-color: #f8f9fa;
        color: #003F87;
        font-weight: 500;
    }
    
    /* Success messages */
    .success {
        color: #0075C9;
    }
    
    /* Download button styling */
    .stDownloadButton > button {
        background-color: #0075C9;
        color: white;
        border: none;
    }
    
    .stDownloadButton > button:hover {
        background-color: #003F87;
    }
    
    /* Tab styling */
    .stTabs [data-baseweb="tab-list"] {
        gap: 24px;
    }
    
    .stTabs [data-baseweb="tab"] {
        font-weight: 500;
        color: #003F87;
    }
    
    .stTabs [aria-selected="true"] {
        color: #0075C9;
    }
    
    /* Progress info box */
    .progress-info {
        background-color: #f0f8ff;
        border: 1px solid #0075C9;
        border-radius: 5px;
        padding: 10px;
        margin: 10px 0;
    }
</style>
""", unsafe_allow_html=True)

st.title("Program Process Analyzer & Overlap Detector")
st.markdown("**Tyler Technologies** - Government Process Optimization Tool")
st.markdown("---")

# Initialize session state
if 'analysis_complete' not in st.session_state:
    st.session_state.analysis_complete = False
if 'program_processes' not in st.session_state:
    st.session_state.program_processes = {}
if 'overlap_analysis' not in st.session_state:
    st.session_state.overlap_analysis = ""
if 'analyses' not in st.session_state:
    st.session_state.analyses = []
if 'process_overlap_data' not in st.session_state:
    st.session_state.process_overlap_data = {}
if 'current_batch' not in st.session_state:
    st.session_state.current_batch = 0
if 'processing_complete' not in st.session_state:
    st.session_state.processing_complete = False

uploaded_file = st.file_uploader("Upload Excel Spreadsheet of Programs", type=['xlsx'])

if uploaded_file:
    df = pd.read_excel(uploaded_file)
    st.success("File uploaded successfully!")
    st.dataframe(df)
    
    total_programs = len(df)
    st.info(f"Total programs to analyze: {total_programs}")
    
    # Check for saved progress
    saved_progress = load_progress()
    if saved_progress and not st.session_state.processing_complete:
        st.warning("Previous analysis was interrupted. Would you like to resume?")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Resume Previous Analysis"):
                st.session_state.program_processes = saved_progress.get('program_processes', {})
                st.session_state.current_batch = saved_progress.get('current_batch', 0)
                st.session_state.analyses = saved_progress.get('analyses', [])
        with col2:
            if st.button("Start Fresh"):
                st.session_state.program_processes = {}
                st.session_state.current_batch = 0
                st.session_state.analyses = []
                st.session_state.saved_progress = None

    if st.button("Analyze Programs and Find Process Overlaps") or (saved_progress and st.session_state.current_batch > 0):
        # Start timing
        start_time = time.time()
        
        # Batch processing setup
        total_batches = (total_programs + BATCH_SIZE - 1) // BATCH_SIZE
        
        # Progress tracking
        progress_container = st.container()
        progress_text = st.empty()
        progress_bar = st.progress(0)
        time_info = st.empty()
        
        # Phase 1: Process programs in batches
        st.header("Phase 1: Identifying Key Processes")
        
        # Process from current batch to end
        for batch_num in range(st.session_state.current_batch, total_batches):
            batch_start = batch_num * BATCH_SIZE
            batch_end = min(batch_start + BATCH_SIZE, total_programs)
            
            st.subheader(f"Processing Batch {batch_num + 1} of {total_batches}")
            batch_info = st.info(f"Programs {batch_start + 1} to {batch_end} of {total_programs}")
            
            # Process programs in this batch
            for idx in range(batch_start, batch_end):
                row = df.iloc[idx]
                program_name = row['Program']
                program_description = row['Description']
                
                # Skip if already processed
                if program_name in st.session_state.program_processes:
                    continue
                
                # Update progress
                overall_progress = (idx + 1) / total_programs
                progress_bar.progress(overall_progress)
                progress_text.text(f"Processing '{program_name}' ({idx + 1}/{total_programs})")
                
                # Update time estimate
                elapsed = time.time() - start_time
                time_remaining = estimate_time_remaining(idx - batch_start + 1, total_programs - batch_start, elapsed)
                time_info.text(f"Elapsed: {str(timedelta(seconds=int(elapsed)))} | Estimated remaining: {time_remaining}")
                
                # Process program
                process_prompt = f"""
                Analyze the following government program and identify its KEY PROCESSES.

                Program: {program_name}
                Description: {program_description}

                List the major processes within this program. For each process:
                1. Give it a clear, concise name
                2. Provide a brief description (1-2 sentences)
                3. Identify the process type (e.g., process types like inspections, enforcement, auditing, permitting and licensing, investigations, application intake, application review, payment handling, procurement and purchasing, records management, customer service, public notification, public engagement, benefits distribution, grant administration, data collection and reporting, strategic planning, budget development, performance management, policy and ordinance development, etc.)

                Format your response as a JSON array like this:
                [
                    {{
                        "process_name": "Process Name",
                        "description": "Brief description",
                        "process_type": "Type of process"
                    }},
                    ...
                ]
                """

                try:
                    response = client.chat.completions.create(
                        model="gpt-4o",
                        messages=[{"role": "user", "content": process_prompt}],
                        temperature=0.3
                    )

                    # Extract JSON from response
                    response_text = response.choices[0].message.content
                    json_match = re.search(r'\[\s*\{.*\}\s*\]', response_text, re.DOTALL)
                    
                    if json_match:
                        json_text = json_match.group()
                        processes = json.loads(json_text)
                    else:
                        processes = json.loads(response_text)
                    
                    st.session_state.program_processes[program_name] = processes
                    
                    # Display processes
                    with st.expander(f"Processes for {program_name}"):
                        for proc in processes:
                            st.write(f"**{proc['process_name']}** ({proc['process_type']})")
                            st.write(f"   {proc['description']}")
                    
                except Exception as e:
                    st.warning(f"Error processing {program_name}: {str(e)}")
                    st.session_state.program_processes[program_name] = [{
                        "process_name": "Process extraction failed",
                        "description": "Unable to parse structured process data.",
                        "process_type": "Unknown"
                    }]
                
                # Save progress periodically
                if (idx + 1) % SAVE_INTERVAL == 0:
                    save_progress({
                        'program_processes': st.session_state.program_processes,
                        'current_batch': batch_num,
                        'analyses': st.session_state.analyses
                    })
                    st.info(f"Progress saved at program {idx + 1}")
            
            # Update current batch
            st.session_state.current_batch = batch_num + 1
            batch_info.empty()
        
        st.success("Process identification complete!")
        
        # Phase 2: Overlap Analysis
        st.header("Phase 2: Analyzing Process Overlaps")
        
        with st.spinner("Analyzing process overlaps across all programs..."):
            # Create process overlap analysis data
            process_overlap_data = analyze_process_overlaps(st.session_state.program_processes)
            st.session_state.process_overlap_data = process_overlap_data
        
        # Visualizations
        st.header("Process Overlap Visualizations")
        
        # Create tabs for different visualizations
        tab1, tab2, tab3, tab4 = st.tabs(["Overlap Matrix", "Process Type Distribution", "Network Graph", "Cross-Department Analysis"])
        
        with tab1:
            create_overlap_matrix(process_overlap_data, st.session_state.program_processes)
        
        with tab2:
            create_process_type_chart(st.session_state.program_processes)
            
        with tab3:
            create_network_graph(process_overlap_data, st.session_state.program_processes)
            
        with tab4:
            create_cross_department_analysis(process_overlap_data, st.session_state.program_processes)
        
        # Generate overlap analysis text
        all_processes_text = ""
        for program, processes in st.session_state.program_processes.items():
            all_processes_text += f"\n\nProgram: {program}\n"
            for proc in processes:
                all_processes_text += f"- {proc['process_name']} ({proc['process_type']}): {proc['description']}\n"

        # Truncate if too long for API
        if len(all_processes_text) > 50000:
            all_processes_text = all_processes_text[:50000] + "\n\n[Content truncated due to length...]"

        overlap_prompt = f"""
        Analyze the following processes from multiple government programs and identify opportunities for consolidation or centralization.

        {all_processes_text}

        Provide a comprehensive analysis including:

        1. **Common Process Types**: Identify process types that appear across multiple programs
        2. **Specific Process Overlaps**: List specific processes that are similar or identical across programs
        3. **Consolidation Opportunities**: Describe specific opportunities to consolidate or centralize processes
        4. **Shared Service Recommendations**: Suggest which processes could become shared services
        5. **Implementation Priority**: Rank consolidation opportunities by potential impact and ease of implementation
        6. **Expected Benefits**: Quantify potential efficiency gains, cost savings, or service improvements

        Be specific about which programs share which processes and how they could be consolidated.
        """

        with st.spinner("Generating comprehensive overlap analysis..."):
            overlap_response = client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": overlap_prompt}],
                temperature=0.5
            )
            
            overlap_analysis = overlap_response.choices[0].message.content
            st.session_state.overlap_analysis = overlap_analysis
        
        st.header("Process Overlap Analysis")
        st.markdown(overlap_analysis)
        
        # Phase 3: Individual analyses in batches
        st.header("Phase 3: Individual Program Efficiency Analyses")
        
        analysis_progress = st.progress(0)
        analysis_text = st.empty()
        
        # Process analyses in batches
        for batch_num in range(0, total_batches):
            batch_start = batch_num * BATCH_SIZE
            batch_end = min(batch_start + BATCH_SIZE, total_programs)
            
            st.subheader(f"Analyzing Batch {batch_num + 1} of {total_batches}")
            
            for idx in range(batch_start, batch_end):
                row = df.iloc[idx]
                program_name = row['Program']
                
                # Skip if already analyzed
                if any(a['Program'] == program_name for a in st.session_state.analyses):
                    continue
                
                program_description = row['Description']
                
                analysis_progress.progress((idx + 1) / total_programs)
                analysis_text.text(f"Analyzing '{program_name}' ({idx + 1}/{total_programs})")
                
                # Get processes for this program
                processes_text = ""
                if program_name in st.session_state.program_processes:
                    processes_text = "\n".join([f"- {p['process_name']}: {p['description']}" 
                                               for p in st.session_state.program_processes[program_name]])

                analysis_prompt = f"""
                You're analyzing a local government program called '{program_name}' for efficiency and cost savings.

                Program description:
                {program_description}

                Key Processes already identified:
                {processes_text}

                Provide a clearly structured analysis including these sections:

                1. Current State Assessment
                2. Areas to Analyze for Efficiency Opportunities
                3. Key Processes within this Program (use the processes identified above and expand if needed)
                4. Types of Recommendations (Streamlining, Automation, Staffing, Fees, Policy, Customer Service enhancements)
                5. Ideal Efficiency Metrics (suggest measurable metrics for tracking improvements)
                6. Anticipated Outcomes and Benefits

                Clearly organize your response into these sections.
                """

                try:
                    response = client.chat.completions.create(
                        model="gpt-4o",
                        messages=[{"role": "user", "content": analysis_prompt}],
                        temperature=0.5
                    )

                    analysis = response.choices[0].message.content
                    st.session_state.analyses.append({'Program': program_name, 'Analysis': analysis})
                    
                except Exception as e:
                    st.warning(f"Error analyzing {program_name}: {str(e)}")
                    st.session_state.analyses.append({
                        'Program': program_name, 
                        'Analysis': f"Analysis failed: {str(e)}"
                    })
                
                # Save progress
                if (idx + 1) % SAVE_INTERVAL == 0:
                    save_progress({
                        'program_processes': st.session_state.program_processes,
                        'current_batch': st.session_state.current_batch,
                        'analyses': st.session_state.analyses
                    })
        
        st.success("All analyses generated!")
        st.session_state.analysis_complete = True
        st.session_state.processing_complete = True
        
        # Clear saved progress
        if 'saved_progress' in st.session_state:
            del st.session_state.saved_progress

    # Display results if analysis is complete
    if st.session_state.analysis_complete:
        # Display individual analyses
        st.header("Individual Program Analyses")
        
        # Add search/filter capability for large results
        search_term = st.text_input("Search programs:", "")
        filtered_analyses = [a for a in st.session_state.analyses 
                           if search_term.lower() in a['Program'].lower()]
        
        st.info(f"Showing {len(filtered_analyses)} of {len(st.session_state.analyses)} analyses")
        
        for item in filtered_analyses[:50]:  # Show first 50 to avoid UI overload
            with st.expander(f"Analysis for {item['Program']}"):
                st.markdown(item['Analysis'])
        
        if len(filtered_analyses) > 50:
            st.warning(f"Showing first 50 results. Use search to find specific programs or download full report.")

        # Download section
        st.markdown("---")
        st.header("Download Reports")
        
        col1, col2 = st.columns(2)
        
        with col1:
            with st.spinner("Generating Word document..."):
                word_file = create_word_doc(st.session_state.program_processes, 
                                           st.session_state.overlap_analysis, 
                                           st.session_state.analyses, 
                                           st.session_state.process_overlap_data)
                st.download_button("Download Complete Word Report", 
                                 word_file, 
                                 "Program_Process_Analysis_Report.docx",
                                 mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        with col2:
            with st.spinner("Generating Excel document..."):
                excel_file = create_excel_file(st.session_state.program_processes, 
                                             st.session_state.overlap_analysis, 
                                             st.session_state.analyses, 
                                             st.session_state.process_overlap_data)
                st.download_button("Download Complete Excel Report", 
                                 excel_file, 
                                 "Program_Process_Analysis_Report.xlsx",
                                 mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")